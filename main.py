import telebot
import io
import PyPDF2
from docx import Document
from pydub import AudioSegment
import ffmpeg
import easyocr
import os
from docx2pdf import convert
import tempfile
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

bot = telebot.TeleBot('7136593333:AAEoE6Zx-0fCblBsbQFNCtoVa0LnoyLIYfo')
reader = easyocr.Reader(['en', 'ru'])


@bot.message_handler(commands=['start'])
def main(message):
    bot.send_message(message.chat.id, 'Привет, этот бот способен конвертировать файлы.\n'
                                      'Для просмотра команд напишите /menu')


@bot.message_handler(commands=['menu'])
def send_menu(message):
    bot.send_message(message.chat.id, '/photo Команда для составления фото отчета.\n'
                                      '/document Команда для конвертирования файлов(pdf <=> docx)\n')


@bot.message_handler(commands=['photo'])
def main(message):
    bot.send_message(message.chat.id, 'Отправьте этикетку для составления отчета')


@bot.message_handler(commands=['document'])
def main(message):
    bot.send_message(message.chat.id, 'Отправьте файл для конвертации')


@bot.message_handler(content_types=['photo'])
def handle_photo(message):
    with tempfile.TemporaryDirectory() as temp_dir:
        file_info = bot.get_file(message.photo[-1].file_id)
        downloaded_file = bot.download_file(file_info.file_path)

        temp_image_path = os.path.join(temp_dir, "temp_image.jpg")
        with open(temp_image_path, "wb") as temp_image:
            temp_image.write(downloaded_file)

        result = reader.readtext(temp_image_path)

        doc = Document()
        text_lines = [clean_and_format_text(detection[1]) for detection in result]
        combined_lines = combine_lines(text_lines)
        formatted_lines = format_lines(combined_lines)

        for line in formatted_lines:
            paragraph = doc.add_paragraph(line)
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

        temp_doc_path = os.path.join(temp_dir, "temp_doc.docx")
        doc.save(temp_doc_path)

        with open(temp_doc_path, "rb") as doc_file:
            bot.send_document(message.chat.id, doc_file)

        # Теперь конвертируем в формат .prn
        converted_prn = convert_docx_to_prn(temp_doc_path)
        # Отправляем конвертированный документ пользователю как файл с расширением .prn
        with open(converted_prn, "rb") as prn_file:
            bot.send_document(message.chat.id, prn_file, caption="Конвертированный документ.prn")

    # Удаляем временные файлы
    if os.path.exists(temp_image_path):
        os.remove(temp_image_path)
    if os.path.exists(temp_doc_path):
        os.remove(temp_doc_path)
    if os.path.exists(converted_prn):
        os.remove(converted_prn)


@bot.message_handler(content_types=['document'])
def handle_document(message):
    # Получаем информацию о документе
    file_info = bot.get_file(message.document.file_id)
    downloaded_file = bot.download_file(file_info.file_path)

    # Определяем формат документа
    file_extension = message.document.file_name.split('.')[-1]

    if file_extension == 'pdf':
        # Конвертируем PDF в DOCX
        converted_docx = convert_pdf_to_docx(downloaded_file)
        # Сохраняем конвертированный документ во временный файл
        with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as temp_doc:
            temp_doc.write(converted_docx.getvalue())
            temp_docx_path = temp_doc.name
        # Отправляем конвертированный документ пользователю как файл с расширением .docx
        bot.send_document(message.chat.id, open(temp_docx_path, "rb"), caption="Конвертированный документ.docx")
        os.remove(temp_docx_path)

    elif file_extension == 'docx':
        # Сохраняем docx файл во временный файл
        with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as temp_doc:
            temp_doc.write(downloaded_file)
            temp_docx_path = temp_doc.name

        # Конвертируем DOCX в PDF
        converted_pdf = convert_docx_to_pdf(temp_docx_path)
        # Отправляем конвертированный документ пользователю как файл с расширением .pdf
        bot.send_document(message.chat.id, open(converted_pdf, "rb"), caption="Конвертированный документ.pdf")
        os.remove(converted_pdf)
        os.remove(temp_docx_path)

    else:
        bot.reply_to(message, "Поддерживаются только документы в форматах PDF и DOCX")


@bot.message_handler(content_types=['audio'])
def handle_audio(message):
    # Получаем информацию об аудио сообщении
    file_info = bot.get_file(message.audio.file_id)
    downloaded_file = bot.download_file(file_info.file_path)

    # Определяем формат аудио
    file_extension = message.audio.mime_type.split('/')[-1]

    if file_extension == 'wav':
        # Конвертируем аудио из WAV в OGG
        converted_audio = convert_audio(downloaded_file, 'wav', 'ogg')
        bot.send_voice(message.chat.id, converted_audio)
    elif file_extension == 'ogg':
        # Конвертируем аудио из OGG в WAV
        converted_audio = convert_audio(downloaded_file, 'ogg', 'wav')
        bot.send_document(message.chat.id, converted_audio, caption="Конвертированный аудиофайл.wav")
    else:
        bot.reply_to(message, "Поддерживаются только аудиоформаты WAV и OGG")


@bot.message_handler(content_types=['video'])
def handle_video(message):
    # Получаем информацию о видео сообщении
    file_info = bot.get_file(message.video.file_id)
    downloaded_file = bot.download_file(file_info.file_path)

    # Определяем формат видео
    file_extension = message.video.mime_type.split('/')[-1]

    if file_extension == 'mp4':
        # Конвертируем видео из MP4 в WebM
        converted_video = convert_video(downloaded_file, 'mp4', 'webm')
        bot.send_document(message.chat.id, converted_video, caption="Конвертированный видеофайл.webm")
    elif file_extension == 'webm':
        # Конвертируем видео из WebM в MP4
        converted_video = convert_video(downloaded_file, 'webm', 'mp4')
        bot.send_document(message.chat.id, converted_video, caption="Конвертированный видеофайл.mp4")
    else:
        bot.reply_to(message, "Поддерживаются только видеоформаты MP4 и WebM")


def clean_and_format_text(text):
    text = ' '.join(text.split())
    replacements = {
        "email: kos@kos ru": "email: kos@kos.ru",
        "+7(843) 533 _98~09": "+7(843) 533-98-09",
        "e- mail": "email",
        "+7(843) 533 98 09": "+7(843) 533-98-09",
        "Polyethylene 11503 070,": "Polyethylene 11503 - 070,",
        "LOT Ng": "LOT №",
        "Dale": "Date",
        "Pallet Ng": "Pallet - №",
        "Malerial": "Material",
        "Net Weight;": "Net Weight",
        "Code kg": "kg",
        "Material Polyethylene ": "Material == Polyethylene ",
        "Polyelhylene": "Polyethylene",
        "Polycarbonale": "Polycarbonate",
        "Pallel": "Pallet",
    }
    for old, new in replacements.items():
        text = text.replace(old, new)
    return text


def combine_lines(text_lines):
    combined_lines = []
    current_line = ""
    for line in text_lines:
        if any(keyword in line for keyword in ["KAZANORGSINTEZ", "Belomorskaya", "email", "Polyethylene", "LOT", "Date", "Pallet", "Material", "Net Weight", "Code"]):
            if current_line:
                combined_lines.append(current_line.strip())
            current_line = line
        else:
            current_line += " " + line
    if current_line:
        combined_lines.append(current_line.strip())
    return combined_lines


def format_lines(lines):
    formatted_lines = []
    for line in lines:
        if "420051" in line:
            line = line.replace("420051", "420051\n")
        if "LOT Ng" in line:
            line = line.replace("LOT Ng", "LOT № ==")
        if "Date" in line:
            line = line.replace("Date", "Date ==")
        if "Pallet Ng" in line:
            line = line.replace("Pallet Ng", "Pallet № ==")
        if "Pallet _ Ng" in line:
            line = line.replace("Pallet _ Ng", "Pallet № ==")
        if "Material Code" in line:
            line = line.replace("Material Code", "Material Code ==")
        if "Net Weight" in line:
            line = line.replace("Net Weight", "Net Weight, kg ==")
        if "Polyethylene 15813 020" in line:
            line = line.replace("Polyethylene 15813 020", "Polyethylene 15813-020")
        if "Polyethylene 15813  020" in line:
            line = line.replace("Polyethylene 15813  020", "Polyethylene 15813-020")
        if "[01" in line:
            line = line.replace("[01", "101")
        if "kos@kos ru" in line:
            line = line.replace("kos@kos ru", "kos@kos.ru")
        if "e- mail" in line:
            line = line.replace("e- mail", "email")
        if "e mail" in line:
            line = line.replace("e mail", "email")
        if "(TР" in line:
            line = line.replace("(TР", "TP\n")
        if "7(" in line:
            line = line.replace("7(", "+7(")
        if " ~" in line:
            line = line.replace(" ~", "-")
        formatted_lines.append(line)
    return formatted_lines


def convert_audio(audio_bytes, from_format, to_format):
    audio = AudioSegment.from_file(io.BytesIO(audio_bytes), format=from_format)
    output_stream = io.BytesIO()
    audio.export(output_stream, format=to_format)
    output_stream.seek(0)
    return output_stream


def convert_video(video_bytes, from_format, to_format):
    input_stream = io.BytesIO(video_bytes)
    output_stream = io.BytesIO()

    with tempfile.NamedTemporaryFile(delete=False, suffix="." + from_format) as temp_input:
        temp_input.write(input_stream.read())
        temp_input_path = temp_input.name

    output_filename = "temp_output." + to_format

    (
        ffmpeg
        .input(temp_input_path)
        .output(output_filename)
        .run()
    )

    with open(output_filename, "rb") as f:
        output_stream.write(f.read())

    output_stream.seek(0)
    os.remove(temp_input_path)
    os.remove(output_filename)
    return output_stream


def convert_pdf_to_docx(pdf_bytes):
    pdf_reader = PyPDF2.PdfReader(io.BytesIO(pdf_bytes))  # Замена PdfFileReader на PdfReader
    doc = Document()

    for page_num in range(len(pdf_reader.pages)):
        page = pdf_reader.pages[page_num]
        doc.add_paragraph(page.extract_text())

    output_stream = io.BytesIO()
    doc.save(output_stream)
    output_stream.seek(0)
    return output_stream


def convert_docx_to_pdf(docx_path):
    # Конвертируем DOCX в PDF используя docx2pdf
    pdf_path = docx_path.replace(".docx", ".pdf")
    convert(docx_path, pdf_path)
    return pdf_path


def convert_docx_to_prn(docx_path):
    # Открываем файл DOCX
    doc = Document(docx_path)

    # Извлекаем текст из документа
    text_content = ""
    for paragraph in doc.paragraphs:
        text_content += paragraph.text + "\n"

    # Создаем файл PRN и записываем в него текст
    prn_path = docx_path.replace(".docx", ".prn")
    with open(prn_path, "w", encoding="utf-8") as prn_file:
        prn_file.write(text_content)

    return prn_path


bot.polling(none_stop=True)
